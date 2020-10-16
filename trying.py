from docx import Document
from docx.shared import Inches

doc = Document()
r1 = doc.add_paragraph("hello")
run = r1.add_run("hi")
run.bold = True
f = run.font
f.highlight_color = 7

doc.save("test.docx")